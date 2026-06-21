import os
import re
import logging
import yaml
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Inches, Pt
from docxtpl import DocxTemplate

DEFAULT_LECTURE_CRITERIA = """| Оценка | Критерии оценивания |
|---|---|
| Высокий уровень «5» (отлично) | оценку «отлично» заслуживает студент, который свободно ориентируется в теме занятия, четко аргументирует собственную позицию, ставит уточняющие вопросы, выделяет ключевые проблемы обсуждения, подводит промежуточные итоги, проявляет искреннюю заинтересованность во мнении других участников обсуждения, правильно отвечает на вопросы преподавателя. Компетенции, закреплённые за дисциплиной, сформированы на уровне – высокий. |
| Средний уровень «4» (хорошо) | оценку «хорошо» заслуживает студент, практически полностью освоивший знания, умения, компетенции и теоретический материал, учебные задания не оценены максимальным числом баллов, в основном сформировал практические навыки. Компетенции, закреплённые за дисциплиной, сформированы на уровне – хороший (средний). |
| Пороговый уровень «3» (удовлетворительно) | оценку «удовлетворительно» заслуживает студент, частично с пробелами освоивший знания, умения, компетенции и теоретический материал, многие учебные задания либо не выполнил, либо они оценены числом баллов близким к минимальному, некоторые практические навыки не сформированы. Компетенции, закреплённые за дисциплиной, сформированы на уровне – достаточный. |
| Минимальный уровень «2» (неудовлетворительно) | оценку «неудовлетворительно» заслуживает студент, не освоивший знания, умения, компетенции и теоретический материал, учебные задания не выполнил, практические навыки не сформированы. Компетенции, закреплённые за дисциплиной, не сформированы. """

DEFAULT_PRACTICAL_CRITERIA = """| Оценка | Критерии оценивания |
|---|---|
| Высокий уровень «5» (отлично) | Оценка «пять» ставится в том случае, если студент верно решил задачу, свободно владеет методикой расчета, обладает необходимыми теоретическими знаниями, выводы обоснованы. Компетенции, закреплённые за дисциплиной, сформированы на уровне – высокий. |
| Средний уровень «4» (хорошо) | Оценка «четыре» ставится в том случае, если во время защиты работы, при верно выполненных расчетах, преподавателю приходилось периодически задавать студенту уточняющие/пояснительные вопросы для выяснения глубины знаний. Компетенции, закреплённые за дисциплиной, сформированы на уровне – хороший (средний). |
| Пороговый уровень «3» (удовлетворительно) | Оценка «три» выставляется студенту, если задача решена неверно, однако после замечания преподавателя студент нашел и исправил ошибку, а также если при верно выполненных расчетах во время защиты работы студент продемонстрировал отрывочные знания теоретической базы и методики выполнения расчетов. Компетенции, закреплённые за дисциплиной, сформированы на уровне – достаточный. |
| Минимальный уровень «2» (неудовлетворительно) | Оценка «два» выставляется в том случае, если студент не справился с задачей и не смог объяснить суть работы и ответить положительно ни на один вопрос преподавателя. Компетенции, закреплённые за дисциплиной, не сформированы. """

DEFAULT_CREDIT_CRITERIA = """| Оценка | Критерии оценивания |
|---|---|
| Высокий уровень «5» (отлично) | оценку «отлично» заслуживает студент, освоивший знания, умения, компетенции и теоретический материал без пробелов; выполнивший все задания, предусмотренные учебным планом на высоком качественном уровне; практические навыки профессионального применения освоенных знаний сформированы. |
| Средний уровень «4» (хорошо) | оценку «хорошо» заслуживает студент, практически полностью освоивший знания, умения, компетенции и теоретический материал, учебные задания не оценены максимальным числом баллов, в основном сформировал практические навыки. |
| Пороговый уровень «3» (удовлетворительно) | оценку «удовлетворительно» заслуживает студент, частично с пробелами освоивший знания, умения, компетенции и теоретический материал, многие учебные задания либо не выполнил, либо они оценены числом баллов близким к минимальному, некоторые практические навыки не сформированы. |
| Минимальный уровень «2» (неудовлетворительно) | оценку «неудовлетворительно» заслуживает студент, не освоивший знания, умения, компетенции и теоретический материал, учебные задания не выполнил, практические навыки не сформированы. """

DEFAULT_TEST_PAPER_CRITERIA = """| Шкала оценивания, % правильных ответов от максимально возможного | Оценка |
|---|---|
| 80-100 | Высокий уровень «5» (отлично) |
| 65-79 | Средний уровень «4» (хорошо) |
| 50-64 | Пороговый уровень «3» (удовлетворительно) |
| 49 и менее | Минимальный уровень «2» (неудовлетворительно) """

try:
    from google.adk.tools.tool_context import ToolContext
except ImportError:
    from google.adk.tools import ToolContext

logger = logging.getLogger(__name__)

SECTIONS = {
    "Кейс-задача": "case_study",
    "Вопросы для коллоквиумов, собеседования": "colloquium",
    "Комплект заданий для контрольной работы": "test_paper",
    "Перечень дискуссионных тем для круглого стола": "round_table",
    "Портфолио": "portfolio",
    "Деловая (ролевая) игра": "roleplay",
    "Темы групповых и/или индивидуальных творческих заданий/проектов": "creative_project",
    "Комплект разноуровневых задач/заданий": "multi_level_tasks",
    "Комплект заданий для расчетно-графической работы": "rgr",
    "Темы эссе/рефератов/докладов/сообщений": "essay",
    "Тематика курсовых работ/проектов": "course_work",
    "Вопросы к зачёту/зачёту с оценкой": "credit",
    "Вопросы к экзамену": "exam"
}

def get_safe_filename(num, theme):
    """
    Generates a safe filename based on lesson number and title.
    """
    combined = f"{num}_{theme}".lower()
    combined = re.sub(r'[^a-zA-Z0-9а-яА-ЯёЁ_]', '_', combined)
    combined = re.sub(r'_+', '_', combined)
    return combined.strip('_')[:60] + ".md"

def is_lecture(act_num, eval_tool):
    num_lower = str(act_num).lower()
    tool_lower = str(eval_tool).lower()
    if 'лекция' in num_lower or 'лек.' in num_lower:
        return True
    if 'устный опрос' in tool_lower:
        return True
    return False

def is_practical(act_num, eval_tool):
    num_lower = str(act_num).lower()
    tool_lower = str(eval_tool).lower()
    if 'практическ' in num_lower or 'практ.' in num_lower or 'семинар' in num_lower or 'лаб.' in num_lower or 'лабораторн' in num_lower:
        return True
    if 'решение задач' in tool_lower:
        return True
    return False

def clean_hyphenation(text: str) -> str:
    if not isinstance(text, str):
        return text
    text = text.replace('\u00ad', '')
    text = re.sub(r'([a-zA-Zа-яА-ЯёЁ]+)-\s+([a-zA-Zа-яА-ЯёЁ]+)', r'\1\2', text)
    text = re.sub(r'([a-zA-Zа-яА-ЯёЁ]+)-\s*\n\s*([a-zA-Zа-яА-ЯёЁ]+)', r'\1\2', text)
    text = re.sub(r'\s{2,}', ' ', text)
    return text

def clean_dict_hyphenations(obj):
    if isinstance(obj, dict):
        return {k: clean_dict_hyphenations(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_dict_hyphenations(x) for x in obj]
    elif isinstance(obj, str):
        return clean_hyphenation(obj)
    return obj

def strip_question_numbering(text: str) -> str:
    if not isinstance(text, str):
        return text
    prev = ""
    while prev != text:
        prev = text
        # Remove leading numbers with dots/parentheses/spaces/dashes or markdown list marks (-, *, +)
        text = re.sub(r'^(\d+[\.\)\s-]*|[-*+]\s*)', '', text).strip()
    return text

def clean_all_list_prefixes(data):
    if isinstance(data, dict):
        cleaned = {}
        for k, v in data.items():
            if k in ["questions", "tasks", "topics", "structure", "roles", "reproductive", "reconstructive", "creative", "group_projects", "individual_projects"]:
                if isinstance(v, list):
                    cleaned[k] = [strip_question_numbering(x) for x in v]
                else:
                    cleaned[k] = clean_all_list_prefixes(v)
            else:
                cleaned[k] = clean_all_list_prefixes(v)
        return cleaned
    elif isinstance(data, list):
        return [clean_all_list_prefixes(x) for x in data]
    else:
        return data

def clean_question_prefixes(data):
    return clean_all_list_prefixes(data)

def write_okf_file(file_path: str, metadata: dict, content: str):
    """
    Writes a Markdown file with a YAML frontmatter header.
    """
    lines = ["---"]
    for k, v in metadata.items():
        if isinstance(v, list):
            items_str = ", ".join(f"'{x}'" if isinstance(x, str) else str(x) for x in v)
            lines.append(f"{k}: [{items_str}]")
        else:
            escaped_val = str(v).replace('"', '\"')
            lines.append(f"{k}: \"{escaped_val}\"")
    lines.append("---")
    lines.append(content)
    
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

def read_okf_file(file_path: str) -> tuple:
    """
    Reads a Markdown file, parsing its YAML frontmatter header and returning
    a tuple of (metadata_dict, markdown_content_str).
    """
    if not os.path.exists(file_path):
        return {}, ""
        
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, content
        
    closing_idx = -1
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            closing_idx = idx
            break
            
    if closing_idx == -1:
        return {}, content
        
    metadata = {}
    for idx in range(1, closing_idx):
        line = lines[idx].strip()
        if not line or ":" not in line:
            continue
        parts = line.split(":", 1)
        k = parts[0].strip()
        v = parts[1].strip()
        
        if v.startswith("[") and v.endswith("]"):
            items = v[1:-1].split(",")
            parsed_items = []
            for item in items:
                item_strip = item.strip()
                if (item_strip.startswith("'") and item_strip.endswith("'")) or (item_strip.startswith('"') and item_strip.endswith('"')):
                    parsed_items.append(item_strip[1:-1])
                else:
                    parsed_items.append(item_strip)
            metadata[k] = parsed_items
        elif (v.startswith('"') and v.endswith('"')) or (v.startswith("'") and v.endswith("'")):
            metadata[k] = v[1:-1].replace('\\"', '"')
        else:
            metadata[k] = v
            
    markdown_content = "\n".join(lines[closing_idx+1:])
    return metadata, markdown_content

def parse_markdown_table_file(file_path: str) -> list:
    _, content = read_okf_file(file_path)
    if not content:
        return []
    
    rows = []
    for line in content.strip().split("\n"):
        line = line.strip()
        if not line.startswith("|") or not line.endswith("|"):
            continue
        parts = line.split("|")
        parts = parts[1:-1]
        parts = [p.strip().replace("\\|", "|") for p in parts]
        if all(c in "- :" for c in "".join(parts)):
            continue
        rows.append(parts)
    return rows

def insert_table_at_paragraph(doc, paragraph, table_rows):
    if not table_rows:
        paragraph._element.getparent().remove(paragraph._element)
        return
    
    num_rows = len(table_rows)
    num_cols = len(table_rows[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(table_rows):
        for c_idx, cell_value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_value
            
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                
                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                else:
                    if c_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
                    r.font.color.rgb = None
                    r.font.highlight_color = None

    parent = paragraph._element.getparent()
    parent.insert(parent.index(paragraph._element), table._tbl)
    paragraph._element.getparent().remove(paragraph._element)

def apply_visual_styling(doc_path):
    """
    Enforces the exact layout and typography styles of OMD_example_readonly.docx.
    - Margins: Top 0.59", Bottom 0.79", Left 0.98", Right 0.59"
    - Typography: Times New Roman, 14pt justified body with 1.15 spacing, 12pt table text
    - Removes formatting artifacts (green highlights and blue text colors)
    """
    logger.info(f"Applying visual style constraints from OMD_example_readonly.docx to {doc_path}...")
    doc = Document(doc_path)
    
    # 1. Page Margins (Top 0.59", Bottom 0.79", Left 0.98", Right 0.59")
    for sec in doc.sections:
        sec.top_margin = Inches(0.59)
        sec.bottom_margin = Inches(0.79)
        sec.left_margin = Inches(0.98)
        sec.right_margin = Inches(0.59)
        
    # 2. Format Body Paragraphs
    for p in doc.paragraphs:
        text_strip = p.text.strip()
        if not text_strip:
            continue
            
        # Check if activity type header or individual activity
        is_activity_type_header = text_strip.startswith("Вопросы к ") and text_strip.endswith(":")
        is_individual_activity = bool(re.match(r'^(Лекция|Практическое занятие|Лабораторная работа|Лабораторное занятие|Семинар|Лекцион\w*|Практическ\w*)\b', text_strip, re.IGNORECASE))
        
        if is_activity_type_header or is_individual_activity:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)
            is_heading = True
        else:
            # Detect if paragraph represents a heading, list index, or page title
            is_heading = (
                p.style.name.startswith("Heading") or 
                any(r.font.bold for r in p.runs) or 
                p.text.isupper() or
                len(text_strip) < 50 and text_strip.endswith(":")
            )
        
        if not is_heading:
            p.paragraph_format.line_spacing = 1.15
            # Keep cover page/special alignments unchanged, justify standard left text
            if p.alignment is None or p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            
        for r in p.runs:
            r.font.name = "Times New Roman"
            if not is_heading:
                r.font.size = Pt(14)
            # Remove blue text colors
            if r.font.color and r.font.color.rgb:
                r.font.color.rgb = None
            # Remove green highlights
            r.font.highlight_color = None
            
    # 3. Format Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(12)
                        if r.font.color and r.font.color.rgb:
                            r.font.color.rgb = None
                        r.font.highlight_color = None
                        
    doc.save(doc_path)
    logger.info("Visual styling normalization successfully applied!")

def determine_discipline_or_module(rpd_content: str) -> str:
    content_lower = rpd_content.lower()
    # Check first 3000 characters for mentions of "программа модуля" or "рабочая программа модуля"
    header_part = content_lower[:3000]
    if "программа модуля" in header_part or "программы модуля" in header_part or "рабочая программа модуля" in header_part:
        return "МОДУЛЯ"
    
    # Also check if "модул" is in the first 1000 characters and "дисциплин" is not
    first_1000 = content_lower[:1000]
    if "модул" in first_1000 and "дисциплин" not in first_1000:
        return "МОДУЛЯ"
        
    return "ДИСЦИПЛИНЫ"

def docx_generator_tool(tool_context: ToolContext, template_path: str = "templates/OMD_template.docx", output_path: str = "FTD.01_R_OMD_Generated.docx") -> dict:
    """
    Publisher agent tool: compiles results by saving intermediate files, loading
    user-editable questions from markdown, and generating the styled Word report.
    """
    logger.info("Starting docx generator tool...")
    data = clean_dict_hyphenations(tool_context.state.get("omd_json_data", {}))
    data = clean_question_prefixes(data)
    if not data:
        msg = "No extracted JSON data found in session state['omd_json_data']"
        logger.error(msg)
        return {"status": "error", "message": msg}
        
    try:
        # Retrieve target parameters and regenerate flag from session state
        project_name = tool_context.state.get("project_name", "unknown")
        template_path = tool_context.state.get("template_path", template_path)
        output_path = tool_context.state.get("output_path", output_path)
        regenerate = tool_context.state.get("regenerate", False)
        
        processing_dir = os.path.join("processing", project_name)
        results_dir = os.path.join("results", project_name)
        os.makedirs(processing_dir, exist_ok=True)
        os.makedirs(results_dir, exist_ok=True)
        
        # 1. Export competencies to processing/{project_name}/competencies.md (OKF compliant)
        competencies_path = os.path.join(processing_dir, "competencies.md")
        table2_data = data.get("table2", [])
        if table2_data and not (os.path.exists(competencies_path) and not regenerate):
            try:
                lines = []
                lines.append("# Список формируемых компетенций\n")
                lines.append("| № п/п | Код компетенции | Содержание компетенции (или её части) | Индикаторы компетенций | Знать | Уметь | Владеть |")
                lines.append("|---|---|---|---|---|---|---|")
                for row in table2_data:
                    def get_val(obj, key):
                        if isinstance(obj, dict):
                            return obj.get(key, "")
                        else:
                            return getattr(obj, key, "") or ""
                    
                    num = get_val(row, "num")
                    comp_code = get_val(row, "comp_code")
                    content = get_val(row, "content")
                    indicators = get_val(row, "indicators")
                    know = get_val(row, "know")
                    umeti = get_val(row, "umeti")
                    vladeti = get_val(row, "vladeti")
                    
                    def clean_md_cell(text):
                        return str(text).replace("\n", "<br>").replace("|", "\\|").strip()
                        
                    lines.append(f"| {clean_md_cell(num)} | {clean_md_cell(comp_code)} | {clean_md_cell(content)} | {clean_md_cell(indicators)} | {clean_md_cell(know)} | {clean_md_cell(umeti)} | {clean_md_cell(vladeti)} |")
                
                content_str = "\n".join(lines)
                metadata = {
                    "type": "competency_matrix",
                    "title": "Список формируемых компетенций",
                    "description": f"Матрица формируемых компетенций по дисциплине {data.get('course_title', '')}",
                    "tags": ["syllabus", "competencies"]
                }
                write_okf_file(competencies_path, metadata, content_str)
                logger.info(f"Saved competencies markdown table to {competencies_path}")
            except Exception as ex:
                logger.warning(f"Could not save competencies markdown table: {ex}")

        # 2. Export activities to processing/{project_name}/activities.md (OKF compliant)
        activities_path = os.path.join(processing_dir, "activities.md")
        activities_data = data.get("activities", [])
        if activities_data and not (os.path.exists(activities_path) and not regenerate):
            try:
                lines = []
                lines.append("# Тематический план занятий (Таблица 4)\n")
                lines.append("| № п/п | Название раздела, темы | № и название лекций / практических занятий | Формируемые компетенции | Вид контрольного мероприятия | Кол-во часов |")
                lines.append("|---|---|---|---|---|---|")
                for act in activities_data:
                    def get_val(obj, key):
                        if isinstance(obj, dict):
                            return obj.get(key, "")
                        else:
                            return getattr(obj, key, "") or ""
                    num = get_val(act, "num")
                    theme = get_val(act, "theme")
                    comp_code = get_val(act, "comp_code")
                    eval_tool = get_val(act, "eval_tool")
                    hours = get_val(act, "hours")
                    
                    def clean_md_cell(text):
                        return str(text).replace("\n", "<br>").replace("|", "\\|").strip()
                    
                    lesson_title = num
                    if theme and theme not in num:
                        lesson_title = f"{num}. {theme}"
                    lines.append(f"| {clean_md_cell(num)} | {clean_md_cell(theme)} | {clean_md_cell(lesson_title)} | {clean_md_cell(comp_code)} | {clean_md_cell(eval_tool)} | {clean_md_cell(hours)} |")
                
                content_str = "\n".join(lines)
                metadata = {
                    "type": "curriculum_schedule",
                    "title": "Тематический план занятий (Таблица 4)",
                    "description": f"Календарно-тематическое планирование лекционных и практических занятий курса {data.get('course_title', '')}",
                    "tags": ["syllabus", "activities", "schedule"]
                }
                write_okf_file(activities_path, metadata, content_str)
                logger.info(f"Saved activities markdown table to {activities_path}")
            except Exception as ex:
                logger.warning(f"Could not save activities markdown table: {ex}")

        # 3. Export individual lesson/activity questions to separate markdown files (OKF compliant)
        tables_dir = os.path.join(processing_dir, "assessment_tables")
        os.makedirs(tables_dir, exist_ok=True)
        
        if regenerate:
            # Clean folder first
            for fname in os.listdir(tables_dir):
                fpath = os.path.join(tables_dir, fname)
                if os.path.isfile(fpath) and fpath.endswith(".md"):
                    os.unlink(fpath)
                
        for act in activities_data:
            def get_val(obj, key):
                if isinstance(obj, dict):
                    return obj.get(key, "")
                else:
                    return getattr(obj, key, "") or ""
            num = get_val(act, "num")
            theme = get_val(act, "theme")
            questions = get_val(act, "questions") or []
            if not isinstance(questions, list):
                questions = [questions] if questions else []
                
            safe_name = get_safe_filename(num, theme)
            file_path = os.path.join(tables_dir, safe_name)
            
            if not (os.path.exists(file_path) and not regenerate):
                lines = []
                lines.append(f"# Вопросы для {num}\n")
                lines.append(f"**Тема:** {theme}\n")
                lines.append("## Вопросы / Задания:")
                for idx, q in enumerate(questions):
                    lines.append(f"{idx+1}. {q}")
                    
                content_str = "\n".join(lines)
                metadata = {
                    "type": "assessment_questions",
                    "title": f"Вопросы для {num}",
                    "description": f"Вопросы для самоподготовки и проведения занятий по теме: {theme}",
                    "resource": f"processing/{project_name}/assessment_tables/{safe_name}",
                    "tags": ["assessment", "questions", "lesson"]
                }
                write_okf_file(file_path, metadata, content_str)
        logger.info(f"Saved questions for activities to {tables_dir}")

        # 4. Final Compile stage: Load questions back from assessment_tables/ directory
        # Also, initialize the criteria_tables folder and write default markdown files (OKF compliant)
        criteria_dir = os.path.join(processing_dir, "criteria_tables")
        os.makedirs(criteria_dir, exist_ok=True)
        
        lecture_criteria_path = os.path.join(criteria_dir, "lecture.md")
        if not os.path.exists(lecture_criteria_path) or regenerate:
            metadata = {
                "type": "assessment_criteria",
                "title": "Критерии оценивания (lecture)",
                "description": "Шкалы и критерии оценивания по оценочному средству: lecture",
                "resource": f"processing/{project_name}/criteria_tables/lecture.md",
                "tags": ["assessment", "criteria", "lecture"]
            }
            write_okf_file(lecture_criteria_path, metadata, DEFAULT_LECTURE_CRITERIA)
            logger.info(f"Created default lecture criteria at {lecture_criteria_path}")
            
        practical_criteria_path = os.path.join(criteria_dir, "practical.md")
        if not os.path.exists(practical_criteria_path) or regenerate:
            metadata = {
                "type": "assessment_criteria",
                "title": "Критерии оценивания (practical)",
                "description": "Шкалы и критерии оценивания по оценочному средству: practical",
                "resource": f"processing/{project_name}/criteria_tables/practical.md",
                "tags": ["assessment", "criteria", "practical"]
            }
            write_okf_file(practical_criteria_path, metadata, DEFAULT_PRACTICAL_CRITERIA)
            logger.info(f"Created default practical criteria at {practical_criteria_path}")

        CRITERIA_FILES = {
            "credit": DEFAULT_CREDIT_CRITERIA,
            "exam": DEFAULT_CREDIT_CRITERIA,
            "test_paper": DEFAULT_TEST_PAPER_CRITERIA,
            "case_study": DEFAULT_CREDIT_CRITERIA,
            "round_table": DEFAULT_CREDIT_CRITERIA,
            "portfolio": DEFAULT_CREDIT_CRITERIA,
            "roleplay": DEFAULT_CREDIT_CRITERIA,
            "creative_project": DEFAULT_CREDIT_CRITERIA,
            "multi_level_tasks": DEFAULT_CREDIT_CRITERIA,
            "rgr": DEFAULT_CREDIT_CRITERIA,
            "essay": DEFAULT_CREDIT_CRITERIA,
            "course_work": DEFAULT_CREDIT_CRITERIA,
        }
        
        for k, v in CRITERIA_FILES.items():
            criteria_path = os.path.join(criteria_dir, f"{k}.md")
            if not os.path.exists(criteria_path) or regenerate:
                metadata = {
                    "type": "assessment_criteria",
                    "title": f"Критерии оценивания ({k})",
                    "description": f"Шкалы и критерии оценивания по оценочному средству: {k}",
                    "resource": f"processing/{project_name}/criteria_tables/{k}.md",
                    "tags": ["assessment", "criteria", k]
                }
                write_okf_file(criteria_path, metadata, v)
                logger.info(f"Created default criteria table at {criteria_path}")
                
        # Generate reserved files index.md and log.md (OKF compliant)
        from datetime import datetime
        index_path = os.path.join(processing_dir, "index.md")
        if not (os.path.exists(index_path) and not regenerate):
            index_lines = []
            index_lines.append("# Index of Syllabus Knowledge Bundle\n")
            index_lines.append("This directory contains a standardized Open Knowledge Format (OKF) package for RPD conversion artifacts.\n")
            index_lines.append("## Resources:\n")
            index_lines.append(f"- [activities.md](activities.md) (Type: `curriculum_schedule`) - Timetable schedule table.")
            index_lines.append(f"- [competencies.md](competencies.md) (Type: `competency_matrix`) - Competency map requirements.")
            index_lines.append("\n## Criteria Tables:\n")
            for k in sorted(CRITERIA_FILES.keys()):
                index_lines.append(f"- [criteria_tables/{k}.md](criteria_tables/{k}.md) (Type: `assessment_criteria`) - Assessment criteria for '{k}'.")
            index_lines.append("\n## Lesson Question Files:\n")
            for act in activities_data:
                is_dict = isinstance(act, dict)
                num = act.get("num") if is_dict else getattr(act, "num")
                theme = act.get("theme") if is_dict else getattr(act, "theme")
                safe_name = get_safe_filename(num, theme)
                index_lines.append(f"- [assessment_tables/{safe_name}](assessment_tables/{safe_name}) (Type: `assessment_questions`) - Questions for '{num}'.")
                
            index_content = "\n".join(index_lines)
            index_metadata = {
                "type": "bundle_index",
                "title": f"Index of Syllabus Knowledge Bundle: {project_name}",
                "description": "Index manifest cataloging all available parsed syllabus knowledge resources.",
                "timestamp": datetime.utcnow().isoformat() + "Z"
            }
            write_okf_file(index_path, index_metadata, index_content)
        
        log_path = os.path.join(processing_dir, "log.md")
        if not (os.path.exists(log_path) and not regenerate):
            log_lines = []
            log_lines.append("# Changelog of Syllabus Knowledge Bundle\n")
            log_lines.append(f"- **{datetime.utcnow().isoformat() + 'Z'}**: Successfully initialized and compiled OKF knowledge bundle.")
            log_lines.append(f"  - Generated index.md manifest.")
            log_lines.append(f"  - Extracted 2 main metadata matrix tables (activities, competencies).")
            log_lines.append(f"  - Exported {len(activities_data)} lesson-based assessment files.")
            log_lines.append(f"  - Created {len(CRITERIA_FILES) + 2} criteria evaluation tables.")
            
            log_content = "\n".join(log_lines)
            log_metadata = {
                "type": "bundle_log",
                "title": f"Changelog of Syllabus Knowledge Bundle: {project_name}",
                "description": "History of generation, compilation, and modification events for this bundle.",
                "timestamp": datetime.utcnow().isoformat() + "Z"
            }
            write_okf_file(log_path, log_metadata, log_content)

        updated_activities = []
        lectures_sections = []
        practicals_sections = []
        
        for act in activities_data:
            is_dict = isinstance(act, dict)
            num = act.get("num") if is_dict else getattr(act, "num")
            theme = act.get("theme") if is_dict else getattr(act, "theme")
            eval_tool = act.get("eval_tool") if is_dict else getattr(act, "eval_tool")
            act_type = act.get("type") if is_dict else getattr(act, "type")
            
            safe_name = get_safe_filename(num, theme)
            file_path = os.path.join(tables_dir, safe_name)
            questions = []
            
            if os.path.exists(file_path):
                _, file_content = read_okf_file(file_path)
                in_questions_section = False
                for line in file_content.strip().splitlines():
                    line_str = line.strip()
                    if "## Вопросы" in line_str or "## Задания" in line_str:
                        in_questions_section = True
                        continue
                    if in_questions_section and line_str:
                        q_text = strip_question_numbering(line_str)
                        if q_text:
                            questions.append(q_text)
                            
            if not questions:
                orig_qs = act.get("questions") if is_dict else getattr(act, "questions")
                if isinstance(orig_qs, list):
                    questions = orig_qs
                elif orig_qs:
                    questions = [orig_qs]
                    
            # Ensure loaded/extracted questions are fully cleaned of numbering
            questions = [strip_question_numbering(q) for q in questions]
                    
            if is_dict:
                act["questions"] = questions
                updated_activities.append(act)
            else:
                act.questions = questions
                updated_activities.append(act)
                
            lesson_title = num
            if theme and theme not in num:
                lesson_title = f"{num}. {theme}"
                
            sec_data = {
                "section_title": lesson_title,
                "questions": questions
            }
            
            if is_lecture(num, eval_tool) or act_type == "Лекция":
                lectures_sections.append(sec_data)
            elif is_practical(num, eval_tool) or act_type == "Практическое занятие":
                practicals_sections.append(sec_data)
            else:
                if "лек" in num.lower():
                     lectures_sections.append(sec_data)
                else:
                     practicals_sections.append(sec_data)
            
        data["activities"] = updated_activities
        data["colloquium"] = {
            "lectures": lectures_sections,
            "practicals": practicals_sections,
            "criteria": data.get("colloquium", {}).get("criteria", "") if isinstance(data.get("colloquium"), dict) else (getattr(data.get("colloquium"), "criteria", "") or "")
        }

        # 5. Populate and render document
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template path {template_path} not found.")
        
        # Resolve major type and labels
        rpd_content = tool_context.state.get("rpd_content", "")
        degree_type_str = str(data.get("degree_type", "")).lower()
        
        # Determine whether discipline or module
        data["discipline_or_module"] = determine_discipline_or_module(rpd_content)
        
        is_specialist = False
        if "специалист" in degree_type_str or "специальност" in rpd_content.lower():
            is_specialist = True
            
        if is_specialist:
            data["major_type_label"] = "Специальность"
            data["profile_type_label"] = "Специализация"
            data["major_type_label_genitive"] = "специальности"
            data["major_type_label_prepositional"] = "специальности"
            data["degree_qualification"] = "специалист"
            data["degree_qualification_genitive"] = "специалиста"
        else:
            if "магистр" in degree_type_str:
                data["major_type_label"] = "Направление"
                data["profile_type_label"] = "Направленность"
                data["major_type_label_genitive"] = "направления"
                data["major_type_label_prepositional"] = "направлению"
                data["degree_qualification"] = "магистр"
                data["degree_qualification_genitive"] = "магистра"
            else:
                data["major_type_label"] = "Направление"
                data["profile_type_label"] = "Направленность"
                data["major_type_label_genitive"] = "направления"
                data["major_type_label_prepositional"] = "направлению"
                data["degree_qualification"] = "бакалавр"
                data["degree_qualification_genitive"] = "бакалавра"

        if "бакалавр" in degree_type_str:
            data["degree_type"] = "бакалавров"
        elif "магистр" in degree_type_str:
            data["degree_type"] = "магистров"
        elif "специалист" in degree_type_str:
            data["degree_type"] = "специалистов"
        else:
            if data["degree_qualification"] == "магистр":
                data["degree_type"] = "магистров"
            elif data["degree_qualification"] == "специалист":
                data["degree_type"] = "специалистов"
            else:
                data["degree_type"] = "бакалавров"

        credit_data = data.get("credit")
        is_graded = False
        if "оценкой" in rpd_content.lower() or "оценкой" in str(credit_data).lower():
            is_graded = True
            
        if is_graded:
            data["credit_heading_label"] = "Вопросы к зачёту с оценкой"
        else:
            data["credit_heading_label"] = "Вопросы к зачёту"

        # Populate parameters from session state
        data["year_of_study_start"] = tool_context.state.get("year_of_study_start", "2023")
        data["course_type"] = tool_context.state.get("course_type", "профессиональное обучение")
        data["hours"] = tool_context.state.get("hours", "72")
        data["cathedra_name"] = data.get("department", "")
        
        # Set protocol meeting year
        protocol_year = str(data.get("protocol_year", "____"))
        if len(protocol_year) == 2:
            data["cathedra_meeting_year"] = "20" + protocol_year
        else:
            data["cathedra_meeting_year"] = protocol_year
            
        # Render utilizing docxtpl
        doc_tpl = DocxTemplate(template_path)
        for key in SECTIONS.values():
            if key not in data or data[key] is None:
                data[key] = None
        if "protocol_num" not in data: data["protocol_num"] = "__"
        if "protocol_day" not in data: data["protocol_day"] = "__"
        if "protocol_month" not in data: data["protocol_month"] = "___________"
        if "protocol_year" not in data: data["protocol_year"] = "20__"
        
        # Write all dynamic template variables to variables.yml
        variables_path = os.path.join(processing_dir, "variables.yml")
        if not (os.path.exists(variables_path) and not regenerate):
            try:
                with open(variables_path, "w", encoding="utf-8") as f:
                    yaml.dump(data, f, allow_unicode=True, default_flow_style=False)
                logger.info(f"Saved template variables to {variables_path}")
            except Exception as ex:
                logger.warning(f"Could not save template variables to YAML: {ex}")
            
        doc_tpl.render(data)
        doc_tpl.save(output_path)
        
        # 6. In-place XML insertion for criteria tables
        doc = Document(output_path)
        
        markers = []
        for p in doc.paragraphs:
            val = p.text.strip()
            if val.endswith("_CRITERIA_MARKER") or val in ("LECTURE_CRITERIA_MARKER", "PRACTICAL_CRITERIA_MARKER"):
                markers.append((p, val))
                
        for p, marker_val in markers:
            if marker_val == "LECTURE_CRITERIA_MARKER":
                md_path = os.path.join(processing_dir, "criteria_tables", "lecture.md")
            elif marker_val == "PRACTICAL_CRITERIA_MARKER":
                md_path = os.path.join(processing_dir, "criteria_tables", "practical.md")
            else:
                key = marker_val.replace("_CRITERIA_MARKER", "").lower()
                md_path = os.path.join(processing_dir, "criteria_tables", f"{key}.md")
                
            rows = parse_markdown_table_file(md_path)
            insert_table_at_paragraph(doc, p, rows)
            
        doc.save(output_path)
        
        # 7. Apply post-rendering visual styling engine
        apply_visual_styling(output_path)
        
        logger.info(f"Generated final OMD report at {output_path}")
        return {"status": "success", "message": f"Successfully generated report at {output_path}"}
        
    except Exception as e:
        msg = f"Failed to generate Word document: {str(e)}"
        logger.exception(msg)
        return {"status": "error", "message": msg}
